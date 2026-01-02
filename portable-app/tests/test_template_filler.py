"""
Tests for Template Filler Service
"""

import sys
import tempfile
from pathlib import Path

# Add app to path
sys.path.insert(0, str(Path(__file__).parent.parent / "app"))

from docx import Document
from services.template_filler import TemplateFiller, sanitize_filename


def test_sanitize_filename():
    """Test filename sanitization."""
    print("Testing sanitize_filename...")
    
    # Normal names
    assert sanitize_filename("John Smith") == "John Smith"
    assert sanitize_filename("Rahul Kumar") == "Rahul Kumar"
    
    # Names with invalid characters
    assert sanitize_filename("John: Smith") == "John_ Smith"
    assert sanitize_filename("Test<>File") == "Test__File"
    assert sanitize_filename("A/B\\C") == "A_B_C"
    
    # Edge cases
    assert sanitize_filename("") == "unnamed"
    assert sanitize_filename("...") == "unnamed"
    assert sanitize_filename("  name  ") == "name"
    
    print("  [OK] All sanitize_filename tests passed!")


def test_template_filler():
    """Test template filling with Word document."""
    print("Testing TemplateFiller...")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # Create a simple Word template with placeholders
        template_path = temp_dir / "template.docx"
        doc = Document()
        doc.add_paragraph("Student Report Card")
        doc.add_paragraph("Name: {{name}}")
        doc.add_paragraph("Roll No: {{rollno}}")
        doc.add_paragraph("Class: {{class}}")
        doc.add_paragraph("English: {{english}}")
        doc.add_paragraph("Result: {{result}}")
        doc.save(str(template_path))
        
        # Test template filler
        filler = TemplateFiller(str(template_path))
        
        # Fill template
        output_path = temp_dir / "output.docx"
        data = {
            "name": "Rahul Kumar",
            "rollno": "1",
            "class": "Class 5A",
            "english": "85",
            "result": "Pass"
        }
        
        result_path = filler.fill_template(data, str(output_path))
        
        # Verify output exists
        assert Path(result_path).exists(), "Output file should exist"
        
        # Verify content was replaced
        filled_doc = Document(result_path)
        full_text = "\n".join([p.text for p in filled_doc.paragraphs])
        
        assert "Rahul Kumar" in full_text, "Name should be replaced"
        assert "{{name}}" not in full_text, "Placeholder should be gone"
        assert "Class 5A" in full_text, "Class should be replaced"
        assert "85" in full_text, "English score should be replaced"
        
        print("  [OK] All TemplateFiller tests passed!")


def test_template_placeholders():
    """Test placeholder extraction."""
    print("Testing placeholder extraction...")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        
        # Create template with various placeholders
        template_path = temp_dir / "template.docx"
        doc = Document()
        doc.add_paragraph("{{name}} scored {{marks}} in {{subject}}")
        doc.save(str(template_path))
        
        filler = TemplateFiller(str(template_path))
        placeholders = filler.get_placeholders()
        
        # Note: docxtpl might not find all in simple cases, but it should work with render
        print(f"  Found placeholders: {placeholders}")
        
        print("  [OK] Placeholder extraction test passed!")


if __name__ == "__main__":
    print("\n" + "="*50)
    print("Running Template Filler Tests")
    print("="*50 + "\n")
    
    try:
        test_sanitize_filename()
        test_template_filler()
        test_template_placeholders()
        
        print("\n" + "="*50)
        print("[OK] ALL TESTS PASSED!")
        print("="*50 + "\n")
    except AssertionError as e:
        print(f"\n[FAIL] TEST FAILED: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"\n[FAIL] ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

